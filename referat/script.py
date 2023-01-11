
from os import rename
import wikipedia 
import re 
from docx import Document

name =  input("Introdu numele tau :\n")
wikipedia.set_lang("ro")
title = input("despre ce vrei sa fie proiectul tau? :\n")
while True:
    try :
        wiki = wikipedia.page(title)
        break
    except:
       exit(1)
        
        
      
    
 

text = wiki.content

text = re.sub(r'==', '', text)
text = re.sub(r'\n', '\n  ', text)
split = text.split('vezi si ',1)
text = split[0]

#print(text)

document = Document()

paragraph = document.add_heading( title , 0)
paragraph.aligment = document.add_heading(name , 1)
paragraph.aligment = 1

document.add_paragraph( '    ' + text)
paragraph.aligment = 2 
document.save(title + ".docx")




 